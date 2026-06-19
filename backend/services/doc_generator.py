"""DOC 文档生成核心模块"""

import os
import sys
import json
import copy
import subprocess
from datetime import datetime
from docx import Document
from docx.shared import Pt, Cm, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn


def _find_template(template_dir, base_name):
    """在模板目录中查找模板文件，优先 _v2 版本，其次原始版本

    Args:
        template_dir: 模板目录路径
        base_name: 模板基础文件名（不含扩展名），如 "施工组织设计模板"

    Returns:
        找到的模板文件完整路径；未找到则返回 None
    """
    # 优先查找 _v2 版本
    v2_docx_path = os.path.join(template_dir, f"{base_name}_v2.docx")
    if os.path.exists(v2_docx_path):
        return v2_docx_path
    # 其次查找原始 .docx
    docx_path = os.path.join(template_dir, f"{base_name}.docx")
    if os.path.exists(docx_path):
        return docx_path
    # 最后查找 .doc
    doc_path = os.path.join(template_dir, f"{base_name}.doc")
    if os.path.exists(doc_path):
        return doc_path
    return None


class DocGenerator:
    """文档生成器类，负责根据模板生成各类施工文档"""

    def __init__(self, template_dir=None, output_dir=None):
        """初始化文档生成器

        Args:
            template_dir: 模板文件目录，默认为 backend/templates
            output_dir: 输出文件目录，默认为 backend/output
        """
        base_dir = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
        self.template_dir = template_dir or os.path.join(base_dir, "templates")
        self.output_dir = output_dir or os.path.join(base_dir, "output")
        # .doc 转 .docx 的缓存目录
        self.converted_dir = os.path.join(self.template_dir, "_converted_docx")
        os.makedirs(self.output_dir, exist_ok=True)
        os.makedirs(self.converted_dir, exist_ok=True)

    def _convert_doc_to_docx(self, doc_path):
        """将 .doc 文件转换为 .docx 格式

        优先使用 win32com（Windows + Word），回退到 LibreOffice 命令行。
        转换结果缓存到 _converted_docx 目录，避免重复转换。

        Args:
            doc_path: .doc 文件的绝对路径

        Returns:
            转换后的 .docx 文件路径，转换失败则返回 None
        """
        doc_basename = os.path.basename(doc_path)
        docx_basename = os.path.splitext(doc_basename)[0] + ".docx"
        converted_path = os.path.join(self.converted_dir, docx_basename)

        # 如果已存在转换后的文件且比源文件新，直接返回
        if os.path.exists(converted_path):
            src_mtime = os.path.getmtime(doc_path)
            dst_mtime = os.path.getmtime(converted_path)
            if dst_mtime >= src_mtime:
                return converted_path

        # 方式1：使用 win32com（Windows + Word）
        try:
            import win32com.client
            word = win32com.client.Dispatch("Word.Application")
            word.Visible = False
            doc = word.Documents.Open(os.path.abspath(doc_path))
            doc.SaveAs(os.path.abspath(converted_path), FileFormat=16)  # 16 = wdFormatXMLDocument (.docx)
            doc.Close()
            word.Quit()
            return converted_path
        except Exception:
            pass

        # 方式2：使用 LibreOffice 命令行
        try:
            result = subprocess.run(
                [
                    "soffice", "--headless", "--convert-to", "docx",
                    "--outdir", self.converted_dir,
                    os.path.abspath(doc_path),
                ],
                capture_output=True, timeout=60,
            )
            if result.returncode == 0 and os.path.exists(converted_path):
                return converted_path
        except Exception:
            pass

        print(f"警告：无法将 {doc_basename} 转换为 .docx 格式，将使用兜底文档", file=sys.stderr)
        return None

    def _resolve_template(self, template_path):
        """解析模板路径，处理 .doc 格式自动转换

        Args:
            template_path: 模板文件路径

        Returns:
            可供 python-docx 读取的 .docx 文件路径，不存在则返回 None
        """
        if not os.path.exists(template_path):
            return None

        if template_path.endswith(".docx"):
            return template_path

        if template_path.endswith(".doc"):
            return self._convert_doc_to_docx(template_path)

        return None

    def generate_from_template(self, template_path, output_path, data_dict):
        """根据模板生成文档

        使用 python-docx 打开模板，遍历所有段落和表格，替换 {{占位符}} 为实际数据。
        支持 .doc 和 .docx 格式的模板，.doc 会自动转换为 .docx 后处理。
        支持 equipment_table 和 quality_table 占位符的动态表格行插入。

        Args:
            template_path: 模板文件路径（.doc 或 .docx）
            output_path: 输出文件路径
            data_dict: 占位符与实际数据的字典映射

        Returns:
            生成的文件路径
        """
        resolved = self._resolve_template(template_path)
        if resolved:
            # 模板存在（已转换为.docx或原本就是.docx），基于模板生成
            doc = Document(resolved)
            self._replace_placeholders_in_doc(doc, data_dict)
        else:
            # 模板不存在或转换失败，生成基础格式文档作为兜底
            doc = self._create_fallback_doc(data_dict)

        # 后处理：给所有缺少边框的表格添加边框
        self._ensure_table_borders(doc)

        # 后处理：确保审批表后、第一章前有分页符
        # 场景：模板中"审批表"标题和"一、工程内容"之间只有空段，无分页符，
        # 导致第一章标题排到审批表所在页
        self._ensure_page_break_after_approval(doc)

        # 后处理：处理 data_service 注入的 __PB__ 分页符标记
        # 把"5.1 + 5.2"或"6.1 + 6.2"拆为独立段落，段首插入 <w:br w:type="page"/>
        self._process_page_break_markers(doc)

        # 后处理：移除模板中硬编码的分页符（九、应急处置措施等），避免大片空白
        self._remove_template_page_breaks(doc)

        # 后处理：把含 \n\n 的长段落拆成多个独立段落，改善排版
        self._split_long_paragraphs(doc)

        # 后处理：删除"几乎全空"的表格
        # 场景：模板里有 equipment_table / quality_table 等占位符行，
        # 但 data_dict 没填值时仍会插入 N 行空数据 → 整张表是 100+ 个空行的网格
        # 这里按"非空文本单元格 / 总单元格 < 阈值"判定为"未填充的表"，整张删除
        self._remove_empty_tables(doc)

        # 后处理：删除文档中连续的空段落（模板占位产生的空白区域）
        # 场景：模板里留了大量空 <w:p/> 作占位，生成后这些空段会撑出空白页。
        # 清理中间连续3个以上的空段落（保留1-2个作段落间距），以及末尾连续空段落。
        self._trim_consecutive_empty_paragraphs(doc)

        # 确保输出目录存在
        os.makedirs(os.path.dirname(output_path), exist_ok=True)
        doc.save(output_path)

        # 尝试转换为 PDF
        pdf_path = self._convert_to_pdf(output_path)
        if pdf_path:
            print(f"PDF 已生成: {pdf_path}", file=sys.stderr)

        return output_path, pdf_path

    def _replace_placeholders_in_doc(self, doc, data_dict):
        """替换文档中所有占位符，包括表格行插入

        Args:
            doc: Document 对象
            data_dict: 占位符与实际数据的字典映射
        """
        # 处理表格行插入占位符
        table_placeholders = ["equipment_table", "quality_table"]
        for table in doc.tables:
            self._handle_table_placeholders(table, data_dict, table_placeholders)

        # 调整审批表列宽（签名/意见/日期加宽，会签部门缩窄）
        self._adjust_approval_table_columns(doc)

        # 替换段落中的占位符
        for paragraph in doc.paragraphs:
            self._replace_in_paragraph(paragraph, data_dict)

        # 替换表格单元格中的占位符
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        self._replace_in_paragraph(paragraph, data_dict)

    def _handle_table_placeholders(self, table, data_dict, placeholder_keys):
        """处理表格中的动态行插入占位符

        当表格单元格包含 {{equipment_table}} 或 {{quality_table}} 时，
        删除该占位符行，并在该位置插入数据行。

        Args:
            table: 表格对象
            data_dict: 数据字典
            placeholder_keys: 需要动态插入行的占位符键名列表
        """
        for key in placeholder_keys:
            placeholder = "{{" + key + "}}"
            value = data_dict.get(key)
            if value is None:
                continue

            # 查找包含占位符的行
            target_row_idx = None
            for ri, row in enumerate(table.rows):
                for cell in row.cells:
                    try:
                        cell_text = cell.text
                    except TypeError as e:
                        # python-docx 在某些含空 <w:r> 的 cell 上会因 run.text=None 而崩溃
                        # 这里手动拼接：把 <w:r> 里所有 <w:t> 的文本拼起来
                        from lxml import etree
                        cell_text = ""
                        for t_elem in cell._tc.iter(qn('w:t')):
                            if t_elem.text:
                                cell_text += t_elem.text
                    if placeholder in cell_text:
                        target_row_idx = ri
                        break
                if target_row_idx is not None:
                    break

            if target_row_idx is None:
                continue

            # 给工器具表添加单元格边距（边框由 _ensure_table_borders 统一处理）
            if key == "equipment_table":
                from lxml import etree
                tblPr = table._tbl.find(qn('w:tblPr'))
                if tblPr is not None:
                    # 添加单元格边距
                    for cm in tblPr.findall(qn('w:tblCellMar')):
                        tblPr.remove(cm)
                    tblCellMar = etree.SubElement(tblPr, qn('w:tblCellMar'))
                    for tag, w_val in [('w:top', '850'), ('w:left', '108'),
                                       ('w:bottom', '850'), ('w:right', '108')]:
                        side = etree.SubElement(tblCellMar, qn(tag))
                        side.set(qn('w:w'), w_val)
                        side.set(qn('w:type'), 'dxa')
                    print(f"[doc_generator] equipment_table tblCellMar 已设置: top=850 left=108 bottom=850 right=108 dxa", file=sys.stderr)

            # 解析数据：value 应为列表的列表 [[col1, col2, ...], ...]
            if isinstance(value, str):
                try:
                    rows_data = json.loads(value)
                except (json.JSONDecodeError, TypeError):
                    rows_data = [[value]]
            elif isinstance(value, list):
                rows_data = value
            else:
                rows_data = [[str(value)]]

            if not rows_data:
                continue

            # 获取占位符行的格式（用于复制样式）
            template_row = table.rows[target_row_idx]
            num_cols = len(template_row.cells)

            # 在占位符行之前插入新数据行
            # 使用 python-docx 的方式：在占位符行的 tr 元素之前插入
            insert_before_tr = template_row._tr

            for row_idx, row_data in enumerate(rows_data):
                # 创建新行（复制占位符行的结构）
                new_tr = copy.deepcopy(template_row._tr)
                # 移除行高限制，让行高自动适应内容；
                # 保留 cantSplit：行内不跨页（避免 13 页表格行被"拦腰切两段"）
                from lxml import etree
                trPr = new_tr.find(qn('w:trPr'))
                if trPr is not None:
                    # 移除固定行高（w:trHeight），让行高自动适应内容
                    for trHeight in trPr.findall(qn('w:trHeight')):
                        trPr.remove(trHeight)
                    # 保留 cantSplit（行内不跨页）— 避免 13 页表格行被"拦腰切两段"
                # 设置新行的单元格内容
                for ci, td in enumerate(new_tr.findall(qn('w:tc'))):
                    if ci < len(row_data):
                        # 清除单元格内容并设置新文本
                        for p_elem in td.findall(qn('w:p')):
                            for r_elem in p_elem.findall(qn('w:r')):
                                for t_elem in r_elem.findall(qn('w:t')):
                                    t_elem.text = ""
                        # 在第一个段落中设置文本
                        paragraphs = td.findall(qn('w:p'))
                        if paragraphs:
                            # 若该行是 equipment_table 的第一行，在第一段 run 前插入分页符
                            if key == "equipment_table" and row_idx == 0 and ci == 0:
                                from lxml import etree
                                # 在第一段最前面插入分页符 run
                                # 注意：必须同时加一个空的 <w:t></w:t>，否则 python-docx 的 paragraph.text 会因 run.text=None 而崩溃
                                pb_run = etree.Element(qn('w:r'))
                                pb_br = etree.SubElement(pb_run, qn('w:br'))
                                pb_br.set(qn('w:type'), 'page')
                                pb_t = etree.SubElement(pb_run, qn('w:t'))
                                pb_t.text = ""
                                paragraphs[0].insert(0, pb_run)
                            runs = paragraphs[0].findall(qn('w:r'))
                            if runs:
                                texts = runs[0].findall(qn('w:t'))
                                if texts:
                                    texts[0].text = str(row_data[ci])
                                else:
                                    from lxml import etree
                                    t_elem = etree.SubElement(runs[0], qn('w:t'))
                                    t_elem.text = str(row_data[ci])
                            else:
                                from lxml import etree
                                r_elem = etree.SubElement(paragraphs[0], qn('w:r'))
                                t_elem = etree.SubElement(r_elem, qn('w:t'))
                                t_elem.text = str(row_data[ci])
                    else:
                        # 超出数据列数的单元格清空
                        for p_elem in td.findall(qn('w:p')):
                            for r_elem in p_elem.findall(qn('w:r')):
                                for t_elem in r_elem.findall(qn('w:t')):
                                    t_elem.text = ""

                # 在占位符行之前插入新行
                insert_before_tr.addprevious(new_tr)

            # 删除占位符行
            template_row._tr.getparent().remove(template_row._tr)

            # 移除所有行的固定行高和 cantSplit，允许表格行跨页显示
            # 注意：保持 tblLayout=fixed 不变，autofit 会导致单元格内容跨行时被错误拆分
            for row in table.rows:
                trPr = row._tr.find(qn('w:trPr'))
                if trPr is not None:
                    for trHeight in trPr.findall(qn('w:trHeight')):
                        trPr.remove(trHeight)
                    for cantSplit in trPr.findall(qn('w:cantSplit')):
                        trPr.remove(cantSplit)

    def _remove_empty_tables(self, doc, empty_threshold=0.05):
        """删除"几乎全空"的表格

        判定逻辑：
        - 遍历 doc.body 下所有 w:tbl
        - 计算"非空文本单元格 / 总单元格"比值
        - 若比值 < empty_threshold（默认 5%），视为"未填充的占位符表"，整张删除

        典型场景：施工组织设计模板里的 quality_table 占位符行
          - data_dict 的 quality_control 字段有 117 个 item，但每个 item 没有
            "project / sub_project / basis / method / responsible / record" 字段，
            导致 117 行 × 6 列全空，整张表 99% 都是空单元格
          - 用户反馈"最后一张表格要去掉"
        """
        body = doc.element.body
        # 收集所有 w:tbl 元素（按文档顺序，删除时索引可能变，所以从后往前删）
        tables = body.findall(qn('w:tbl'))
        removed = 0
        for tbl in tables:
            total_cells = 0
            non_empty_cells = 0
            for tc in tbl.findall('.//' + qn('w:tc')):
                total_cells += 1
                # 单元格的文本（含 w:t 元素）
                cell_text = ''.join(tc.itertext()).strip()
                if cell_text:
                    non_empty_cells += 1
            if total_cells == 0:
                continue
            ratio = non_empty_cells / total_cells
            if ratio < empty_threshold:
                tbl.getparent().remove(tbl)
                removed += 1
                print(f'[DocGenerator] 移除空表: 总单元格={total_cells}, 非空={non_empty_cells}, 比例={ratio:.1%}', file=sys.stderr)
        return removed

    def _replace_in_paragraph(self, paragraph, data_dict):
        """替换段落中的占位符

        遍历段落中的所有 run，查找并替换 {{key}} 格式的占位符。

        Args:
            paragraph: 段落对象
            data_dict: 占位符与实际数据的字典映射
        """
        for key, value in data_dict.items():
            # 跳过表格行插入占位符（已由 _handle_table_placeholders 处理）
            if key in ("equipment_table", "quality_table"):
                continue
            placeholder = "{{" + key + "}}"
            if placeholder in paragraph.text:
                # 在 runs 中查找并替换
                for run in paragraph.runs:
                    if placeholder in run.text:
                        run.text = run.text.replace(placeholder, str(value))
                # 如果 runs 中没有完整匹配，尝试整体替换
                if placeholder in paragraph.text:
                    # 合并所有 run 的文本，替换后重新设置
                    full_text = paragraph.text.replace(placeholder, str(value))
                    # 清空所有 run，在第一个 run 中设置完整文本
                    if paragraph.runs:
                        for i, run in enumerate(paragraph.runs):
                            if i == 0:
                                run.text = full_text
                            else:
                                run.text = ""

    def _adjust_approval_table_columns(self, doc):
        """调整审批表列宽：签名/意见/日期加宽，会签部门缩窄

        模板原始列宽（dxa）：监理公司=1099, 会签部门=4543, 签名=1143, 意见=884, 日期=1063
        调整后：监理公司=1099, 会签部门=2200, 签名=1600, 意见=2200, 日期=1633
        总宽不变=8732
        """
        from lxml import etree
        for table in doc.tables:
            # 检查是否是审批表（第一行第一格包含 project_name 或工程名称）
            first_cell_text = ''
            if table.rows:
                first_cell_text = table.rows[0].cells[0].text.strip()[:30]
            # 审批表特征：第2行有"会签部门"
            is_approval = False
            if len(table.rows) > 1:
                row1_texts = [cell.text.strip() for cell in table.rows[1].cells]
                if '会签部门' in row1_texts:
                    is_approval = True
            if not is_approval:
                continue

            tbl = table._tbl
            tblGrid = tbl.find(qn('w:tblGrid'))
            if tblGrid is None:
                continue

            # 新列宽
            new_widths = [1099, 2200, 1600, 2200, 1633]  # 总计 8732
            gridCols = tblGrid.findall(qn('w:gridCol'))
            for ci, col in enumerate(gridCols):
                if ci < len(new_widths):
                    col.set(qn('w:w'), str(new_widths[ci]))

            # 同步更新每行每列的 tcW
            for row in table.rows:
                for ci, cell in enumerate(row.cells):
                    tc = cell._tc
                    tcPr = tc.find(qn('w:tcPr'))
                    if tcPr is None:
                        continue
                    tcW = tcPr.find(qn('w:tcW'))
                    # 跳过合并单元格（gridSpan > 1）
                    gridSpan = tcPr.find(qn('w:gridSpan'))
                    if gridSpan is not None:
                        continue
                    if tcW is not None and ci < len(new_widths):
                        tcW.set(qn('w:w'), str(new_widths[ci]))

            print(f"[doc_generator] 审批表列宽已调整: 监理公司=1099, 会签部门=2200, 签名=1600, 意见=2200, 日期=1633", file=sys.stderr)
            break  # 只处理第一个审批表

    def _ensure_table_borders(self, doc):
        """给所有缺少边框的表格添加标准边框

        模板中的表格通常没有 tblBorders 定义，导致在 Word 中显示为无边框。
        此方法检查每个表格，如果 tblBorders 缺失或所有边框 val 为空，
        则添加标准黑色实线边框。
        """
        from lxml import etree
        for i, table in enumerate(doc.tables):
            tbl = table._tbl
            tblPr = tbl.find(qn('w:tblPr'))
            if tblPr is None:
                tblPr = etree.SubElement(tbl, qn('w:tblPr'))
                tbl.insert(0, tblPr)

            existing_borders = tblPr.find(qn('w:tblBorders'))

            # 判断是否需要添加边框
            need_borders = False
            if existing_borders is None:
                need_borders = True
            else:
                # 检查是否有有效的边框值
                valid_borders = 0
                for child in existing_borders:
                    tag = etree.QName(child).localname
                    val = child.get(qn('w:val'), '')
                    if tag in ('top', 'left', 'bottom', 'right', 'insideH', 'insideV'):
                        if val and val not in ('none', 'nil'):
                            valid_borders += 1
                if valid_borders < 6:
                    need_borders = True

            if not need_borders:
                continue

            # 移除旧的空边框定义
            if existing_borders is not None:
                tblPr.remove(existing_borders)

            # 添加标准边框
            tblBorders = etree.SubElement(tblPr, qn('w:tblBorders'))
            for border_name in ['top', 'left', 'bottom', 'right', 'insideH', 'insideV']:
                border_el = etree.SubElement(tblBorders, qn(f'w:{border_name}'))
                border_el.set(qn('w:val'), 'single')
                border_el.set(qn('w:sz'), '4')
                border_el.set(qn('w:space'), '0')
                border_el.set(qn('w:color'), '000000')

            print(f"[doc_generator] 表格#{i} 添加 tblBorders", file=sys.stderr)

    def _ensure_page_break_after_approval(self, doc):
        """确保封面和审批表分页，审批表和正文分页

        场景：
        1. 模板中"组织技术安全措施审批表"标题在封面页末尾，需要移到第2页
        2. 审批表表格后、"一、工程内容"前需要分页，让正文从新页开始

        处理：
        1. 在审批表标题段落添加 pageBreakBefore，让标题从新页开始
        2. 在审批表表格后、第一个有内容段落添加 pageBreakBefore
        """
        from lxml import etree
        body = doc.element.body
        children = list(body)

        approval_title_idx = None
        approval_table_idx = None

        for i, child in enumerate(children):
            tag = etree.QName(child).localname

            if tag == 'p':
                p_texts = []
                for t_el in child.iter():
                    if etree.QName(t_el).localname == 't' and t_el.text:
                        p_texts.append(t_el.text)
                p_text = ''.join(p_texts).strip()
                if '审批表' in p_text and approval_title_idx is None:
                    approval_title_idx = i

            elif tag == 'tbl' and approval_table_idx is None:
                texts = []
                for t_el in child.iter():
                    if etree.QName(t_el).localname == 't' and t_el.text:
                        texts.append(t_el.text)
                table_text = ''.join(texts)
                if '审批' in table_text or '监理' in table_text or '编制' in table_text:
                    approval_table_idx = i

            if approval_title_idx is not None and approval_table_idx is not None:
                break

        # 1. 审批表标题添加 pageBreakBefore（让标题从新页开始）
        if approval_title_idx is not None:
            title_el = children[approval_title_idx]
            pPr = title_el.find(qn('w:pPr'))
            if pPr is None:
                pPr = etree.SubElement(title_el, qn('w:pPr'))
                title_el.insert(0, pPr)
            if pPr.find(qn('w:pageBreakBefore')) is None:
                etree.SubElement(pPr, qn('w:pageBreakBefore'))
                print(f"[doc_generator] 审批表标题添加 pageBreakBefore", file=sys.stderr)

        # 2. 审批表表格后、第一章前添加 pageBreakBefore
        if approval_table_idx is not None:
            for j in range(approval_table_idx + 1, min(approval_table_idx + 10, len(children))):
                next_child = children[j]
                next_tag = etree.QName(next_child).localname
                if next_tag != 'p':
                    continue
                p_texts = []
                for t_el in next_child.iter():
                    if etree.QName(t_el).localname == 't' and t_el.text:
                        p_texts.append(t_el.text)
                p_text = ''.join(p_texts).strip()

                # 检查是否已有分页符
                has_pb = any(
                    etree.QName(br).localname == 'br'
                    and br.get(qn('w:type')) == 'page'
                    for br in next_child.iter()
                    if etree.QName(br).localname == 'br'
                )
                has_pbb = False
                pPr_check = next_child.find(qn('w:pPr'))
                if pPr_check is not None and pPr_check.find(qn('w:pageBreakBefore')) is not None:
                    has_pbb = True
                if has_pb or has_pbb:
                    break

                if p_text:
                    pPr = next_child.find(qn('w:pPr'))
                    if pPr is None:
                        pPr = etree.SubElement(next_child, qn('w:pPr'))
                        next_child.insert(0, pPr)
                    etree.SubElement(pPr, qn('w:pageBreakBefore'))
                    print(f"[doc_generator] 审批表后正文添加 pageBreakBefore", file=sys.stderr)
                    break

    def _process_page_break_markers(self, doc):
        """处理数据中 __PB__ 分页符标记

        场景：data_service 把内容按"5.1 + 5.2"或"6.1 + 6.2"组合，中间用 \\n\\n__PB__\\n\\n 分隔。
        这里把单个段落按 __PB__ 拆成多段，并在新段落的 run 前面插入 <w:br w:type="page"/>
        （docx 渲染时会在该位置分页）。
        """
        PAGE_BREAK = "__PB__"
        # 找到所有包含 __PB__ 标记的段落（按 body 中顺序）
        body = doc.element.body
        paragraphs = list(doc.paragraphs)
        for p in paragraphs:
            if PAGE_BREAK not in p.text:
                continue
            # 解析该段落：按 __PB__ 切成 N 段；第一段留在原段落，后续段作为新段落插入到原段落之后
            # 段落的所有 run 合并成完整文本，再切分
            full_text = p.text
            # 切分时去掉首尾空白与换行
            raw_segments = full_text.split(PAGE_BREAK)
            # 判断是否以 __PB__ 开头（首段为空白）。这种情况需要在原段落开头加分页符，
            # 否则会被 "去掉空段" 过滤掉，导致首段分页符丢失。
            starts_with_pb = (
                len(raw_segments) > 1
                and raw_segments[0].strip("\n ") == ""
            )
            segments = [s.strip("\n ") for s in raw_segments]
            segments = [s for s in segments if s]  # 去掉空段
            if len(segments) <= 1 and not starts_with_pb:
                continue

            # 复制原段落的 pPr（段落属性）用于新段落
            pPr = p._p.find(qn('w:pPr'))
            pPr_copy_xml = None
            if pPr is not None:
                from lxml import etree
                pPr_copy_xml = etree.tostring(pPr)

            def _add_paragraph_after_spacing(p_elem, after_dxa):
                """给段落 pPr 元素加/更新 w:after 段后空行（单位 dxa）"""
                from lxml import etree
                pPr_local = p_elem.find(qn('w:pPr'))
                if pPr_local is None:
                    pPr_local = etree.SubElement(p_elem, qn('w:pPr'))
                spacing = pPr_local.find(qn('w:spacing'))
                if spacing is None:
                    spacing = etree.SubElement(pPr_local, qn('w:spacing'))
                spacing.set(qn('w:after'), str(after_dxa))

            # 收集原段落所有 run 的格式（取第一个 run 作为模板）
            template_run_xml = None
            if p.runs:
                from lxml import etree
                template_run_xml = etree.tostring(p.runs[0]._r)

            # 清空原段落的所有 run，重写为第一段
            for r in list(p.runs):
                r._r.getparent().remove(r._r)

            # 如果以 __PB__ 开头，先在原段落最前面插入分页符
            from lxml import etree
            if starts_with_pb:
                pb_run = etree.SubElement(p._p, qn('w:r'))
                pb_br = etree.SubElement(pb_run, qn('w:br'))
                pb_br.set(qn('w:type'), 'page')
                pb_t = etree.SubElement(pb_run, qn('w:t'))
                pb_t.text = ""

            # 第一段也按 \n\n 拆成多个子段落（第一个子段落留在原段落，其余插入其后）
            first_sub_parts = [s.strip("\n ") for s in segments[0].split("\n\n")]
            first_sub_parts = [s for s in first_sub_parts if s]
            if first_sub_parts:
                p.add_run(first_sub_parts[0])
            else:
                p.add_run(segments[0])

            # 后续段：每个生成一个新段落，前置 <w:br w:type="page"/>，插入到原段落之后
            from lxml import etree
            current_p_elem = p._p  # 跟踪插入位置
            # 预先从 template_run_xml 提取 rPr（不要复用整段 run，否则会带上原数据里 \n 转成的 <w:br/>）
            template_rPr_xml = None
            if template_run_xml is not None:
                template_root = etree.fromstring(template_run_xml)
                rPr_elem = template_root.find(qn('w:rPr'))
                if rPr_elem is not None:
                    template_rPr_xml = etree.tostring(rPr_elem)
            # 第一段的后续子段落（无分页符，紧跟原段落）
            # 把 \n 也拆成独立段落，让 Word 可以在任意行间分页
            for sub in first_sub_parts[1:]:
                # sub 内可能含 \n，拆成独立行
                sub_lines = [l.strip("\n ") for l in sub.split("\n")]
                sub_lines = [l for l in sub_lines if l]
                for sl in sub_lines:
                    new_p = etree.SubElement(body, qn('w:p'))
                    if pPr_copy_xml is not None:
                        new_pPr = etree.fromstring(pPr_copy_xml)
                        new_p.insert(0, new_pPr)
                    else:
                        etree.SubElement(new_p, qn('w:pPr'))
                    # w:after=0 保持行间距紧凑
                    _add_paragraph_after_spacing(new_p, 0)
                    content_run = etree.SubElement(new_p, qn('w:r'))
                    if template_rPr_xml is not None:
                        content_rPr = etree.fromstring(template_rPr_xml)
                        content_run.append(content_rPr)
                    t_elem = etree.SubElement(content_run, qn('w:t'))
                    t_elem.text = sl
                    t_elem.set('{http://www.w3.org/XML/1998/namespace}space', 'preserve')
                    body.remove(new_p)
                    current_p_elem.addnext(new_p)
                    current_p_elem = new_p
            for seg in segments[1:]:
                # 按 \n\n 拆成多个子段落，再按 \n 拆成独立行
                sub_parts = [s.strip("\n ") for s in seg.split("\n\n")]
                sub_parts = [s for s in sub_parts if s]
                if not sub_parts:
                    continue
                for si, sub in enumerate(sub_parts):
                    # sub 内可能含 \n，拆成独立行
                    sub_lines = [l.strip("\n ") for l in sub.split("\n")]
                    sub_lines = [l for l in sub_lines if l]
                    for sli, sl in enumerate(sub_lines):
                        new_p = etree.SubElement(body, qn('w:p'))
                        if pPr_copy_xml is not None:
                            new_pPr = etree.fromstring(pPr_copy_xml)
                            new_p.insert(0, new_pPr)
                        else:
                            etree.SubElement(new_p, qn('w:pPr'))
                        # w:after=0 保持行间距紧凑
                        _add_paragraph_after_spacing(new_p, 0)
                        # 仅第一个子段落的第一行前加分页符
                        if si == 0 and sli == 0:
                            pb_run = etree.SubElement(new_p, qn('w:r'))
                            pb_br = etree.SubElement(pb_run, qn('w:br'))
                            pb_br.set(qn('w:type'), 'page')
                            pb_t = etree.SubElement(pb_run, qn('w:t'))
                            pb_t.text = ""
                        # 内容 run
                        content_run = etree.SubElement(new_p, qn('w:r'))
                        if template_rPr_xml is not None:
                            content_rPr = etree.fromstring(template_rPr_xml)
                            content_run.append(content_rPr)
                        t_elem = etree.SubElement(content_run, qn('w:t'))
                        t_elem.text = sl
                        t_elem.set('{http://www.w3.org/XML/1998/namespace}space', 'preserve')
                        # 把新段落从 body 末尾移到原段落之后
                        body.remove(new_p)
                        current_p_elem.addnext(new_p)
                        current_p_elem = new_p

    def _remove_template_page_breaks(self, doc):
        """移除模板中硬编码的分页符，让内容自然排版

        场景：模板中"九、应急处置措施"和"10、触电的应急方案"等段落
        自带 <w:br w:type="page"/> 分页符，导致安全措施结束后出现大片空白。
        移除这些分页符，让内容紧跟前文自然排版。
        """
        from lxml import etree
        removed = 0
        for p in doc.paragraphs:
            text = p.text.strip()
            # 只处理特定标题段落中的分页符
            should_remove = False
            if text.startswith('九、') or text.startswith('九、应急处置'):
                should_remove = True
            elif text.startswith('10、') and '触电' in text and '应急' in text:
                should_remove = True

            if not should_remove:
                continue

            # 移除该段落中所有 <w:br w:type="page"/>
            for r in p.runs:
                brs = r._r.findall(qn('w:br'))
                for br in brs:
                    if br.get(qn('w:type')) == 'page':
                        r._r.remove(br)
                        removed += 1
                        print(f"[doc_generator] 移除模板分页符: {text[:40]}", file=sys.stderr)

        if removed:
            print(f"[doc_generator] 共移除模板分页符 {removed} 个", file=sys.stderr)

    def _trim_consecutive_empty_paragraphs(self, doc, max_consecutive=2):
        """删除文档中连续的空段落（包括中间和末尾），保留最多 max_consecutive 个作间距

        场景：模板里留了大量空 <w:p/> 作占位，生成后这些空段会撑出空白页。
        清理策略：
        - 连续3个以上空段落：只保留前 max_consecutive 个
        - 末尾连续空段落：只保留 max_consecutive 个
        - 不删除含分页符的段落
        """
        from lxml import etree
        body = doc.element.body
        children = list(body)
        deleted = 0
        empty_streak = []  # 当前连续空段列表

        def _is_empty_paragraph(el):
            """判断段落是否为空（无文字、无分页符、无图片）"""
            if etree.QName(el).localname != 'p':
                return False
            has_text = False
            for t in el.findall(f".//{qn('w:t')}"):
                if t.text and t.text.strip():
                    has_text = True
                    break
            has_break = bool(el.findall(f".//{qn('w:br')}[@{qn('w:type')}='page']"))
            has_pageBreakBefore = bool(el.findall(f"./{qn('w:pPr')}/{qn('w:pageBreakBefore')}"))
            return not has_text and not has_break and not has_pageBreakBefore

        def _flush_streak():
            nonlocal deleted
            if len(empty_streak) > max_consecutive:
                # 保留前 max_consecutive 个，删除其余
                for el in empty_streak[max_consecutive:]:
                    body.remove(el)
                    deleted += 1
            empty_streak.clear()

        for el in children:
            if _is_empty_paragraph(el):
                empty_streak.append(el)
            else:
                _flush_streak()

        # 处理末尾的连续空段
        _flush_streak()

        if deleted:
            print(f"[doc_generator] 删除连续空段 {deleted} 个（保留每组最多 {max_consecutive} 个）", file=sys.stderr)

    def _trim_trailing_empty_paragraphs(self, doc, max_keep=3):
        """删除文档末尾连续的完全空段落（最多保留 max_keep 段作底部留白）

        场景：模板最后留几十个空 <w:p/> 占位，生成后这些空段会撑出 1-2 个空白页。
        这里从末尾往前连续删除"完全空"的段落（不含文字、不含分页符），
        直到遇到有内容的段落，或已保留 max_keep 段为止。

        注意：<w:pict> / <w:drawing> 视为"可空"——模板末尾的装饰性图片占位不算内容。
        """
        from lxml import etree
        body = doc.element.body
        sectPr = body.find(qn('w:sectPr'))
        deleted = 0
        kept = 0
        children = list(body)
        for el in reversed(children):
            if etree.QName(el).localname != 'p':
                continue
            if el is sectPr:
                continue
            has_text = False
            for t in el.findall(f".//{qn('w:t')}"):
                if t.text and t.text.strip():
                    has_text = True
                    break
            has_break = bool(el.findall(f".//{qn('w:br')}[@{qn('w:type')}='page']"))
            has_pageBreakBefore = bool(el.findall(f"./{qn('w:pPr')}/{qn('w:pageBreakBefore')}"))
            if has_text or has_break or has_pageBreakBefore:
                break  # 遇到真正有内容/分页符就停
            # else: 即使有 pict/drawing 也删除（视为模板末尾占位图）
            if kept < max_keep:
                kept += 1
                continue
            body.remove(el)
            deleted += 1
        if deleted:
            print(f"[doc_generator] 删除末尾空段 {deleted} 个（保留 {kept} 段）", file=sys.stderr)

    def _split_long_paragraphs(self, doc):
        """把含 \\n 或 \\n\\n 的长段落拆成多个独立段落，改善 Word 分页排版

        场景：data_service 把 6.1+6.2 等内容用 \\n\\n 拼接后填入单个占位符，
        导致一个段落里有大量 <w:br/> 换行，Word 无法在段落中间分页。
        拆成独立段落后，Word 可以自然分页。

        关键改动：\\n 换行也拆成独立段落（不再用 <w:br/> 行内换行），
        这样 Word 可以在任意行之间分页，避免文字超出页边距。
        """
        from lxml import etree
        body = doc.element.body
        paragraphs = list(doc.paragraphs)
        for p in paragraphs:
            # 跳过已由 _process_page_break_markers 处理的段落（含 __PB__）
            if "__PB__" in p.text:
                continue
            # 处理含 \n 的段落（\n\n 也包含 \n）
            full_text = p.text
            if "\n" not in full_text:
                continue
            # 按 \n\n 分大段，再按 \n 分小段，全部拆成独立段落
            parts = []
            for segment in full_text.split("\n\n"):
                for line in segment.split("\n"):
                    line = line.strip("\n ")
                    if line:
                        parts.append(line)
            if len(parts) <= 1:
                continue

            # 复制段落属性
            pPr = p._p.find(qn('w:pPr'))
            pPr_copy_xml = None
            if pPr is not None:
                pPr_copy_xml = etree.tostring(pPr)

            # 提取 run 格式
            template_rPr_xml = None
            if p.runs:
                template_root = etree.fromstring(etree.tostring(p.runs[0]._r))
                rPr_elem = template_root.find(qn('w:rPr'))
                if rPr_elem is not None:
                    template_rPr_xml = etree.tostring(rPr_elem)

            # 清空原段落，写入第一部分
            for r in list(p.runs):
                r._r.getparent().remove(r._r)
            p.add_run(parts[0])

            # 后续部分创建新段落（每行一个独立段落，不用 <w:br/>）
            current_p_elem = p._p
            for sub in parts[1:]:
                new_p = etree.SubElement(body, qn('w:p'))
                if pPr_copy_xml is not None:
                    new_pPr = etree.fromstring(pPr_copy_xml)
                    new_p.insert(0, new_pPr)
                else:
                    etree.SubElement(new_p, qn('w:pPr'))
                # 给新段落 pPr 加 w:after=0，保持行间距紧凑
                pPr_local = new_p.find(qn('w:pPr'))
                if pPr_local is not None:
                    spacing = pPr_local.find(qn('w:spacing'))
                    if spacing is None:
                        spacing = etree.SubElement(pPr_local, qn('w:spacing'))
                    spacing.set(qn('w:after'), '0')
                content_run = etree.SubElement(new_p, qn('w:r'))
                if template_rPr_xml is not None:
                    content_rPr = etree.fromstring(template_rPr_xml)
                    content_run.append(content_rPr)
                t_elem = etree.SubElement(content_run, qn('w:t'))
                t_elem.text = sub
                t_elem.set('{http://www.w3.org/XML/1998/namespace}space', 'preserve')
                body.remove(new_p)
                current_p_elem.addnext(new_p)
                current_p_elem = new_p

    def _convert_to_pdf(self, docx_path):
        """将 docx 文件转换为 PDF

        优先使用 LibreOffice 命令行（稳定可靠），
        回退到 docx2pdf（Word COM，在服务进程中可能不稳定）。

        Args:
            docx_path: docx 文件路径

        Returns:
            生成的 PDF 文件路径，失败则返回 None
        """
        pdf_path = os.path.splitext(docx_path)[0] + ".pdf"

        # 如果 PDF 已存在且比 docx 新，直接返回
        if os.path.exists(pdf_path):
            if os.path.getmtime(pdf_path) >= os.path.getmtime(docx_path):
                return pdf_path

        # 方式1：使用 LibreOffice 命令行（优先，稳定可靠）
        try:
            soffice_path = r"C:\Program Files\LibreOffice\program\soffice.exe"
            if not os.path.exists(soffice_path):
                soffice_path = "soffice"  # 回退到 PATH 查找
            result = subprocess.run(
                [
                    soffice_path, "--headless", "--convert-to", "pdf",
                    "--outdir", os.path.dirname(os.path.abspath(docx_path)),
                    os.path.abspath(docx_path),
                ],
                capture_output=True, timeout=120,
                errors='replace',  # 避免中文路径导致的编码错误
            )
            if result.returncode == 0 and os.path.exists(pdf_path):
                print(f"[doc_generator] LibreOffice PDF 转换成功", file=sys.stderr)
                return pdf_path
            else:
                print(f"[doc_generator] LibreOffice 转换失败: rc={result.returncode}", file=sys.stderr)
        except Exception as e:
            print(f"[doc_generator] LibreOffice 异常: {e}", file=sys.stderr)

        # 方式2：使用 docx2pdf（Word COM，在后台服务中可能不稳定）
        try:
            from docx2pdf import convert
            convert(docx_path, pdf_path)
            if os.path.exists(pdf_path):
                print(f"[doc_generator] docx2pdf PDF 转换成功", file=sys.stderr)
                return pdf_path
        except Exception as e:
            print(f"[doc_generator] docx2pdf 异常: {e}", file=sys.stderr)

        return None

    def generate_gantt_chart(self, data_dict, output_path=None):
        """生成 Excel 横道图（与模板格式一致）

        输出格式：
        - 第 1 行：A1 空，B1~末列为月份合并单元格
        - 第 2 行：A1 空，B2~末列为日期序列号（显示为日）
        - A 列：A3~An 为任务名称
        - 横道图条形：对应日期范围的单元格填充蓝色

        Args:
            data_dict: 包含横道图数据的字典，需包含 tasks 列表
                tasks: [{"name": "任务名", "start": "2025-01-01", "end": "2025-01-10", "progress": 50}, ...]
            output_path: 输出文件路径，默认自动生成

        Returns:
            生成的文件路径
        """
        try:
            import openpyxl
            from openpyxl.styles import PatternFill, Font, Alignment, Border, Side
            from openpyxl.utils import get_column_letter
            from openpyxl.worksheet.merge import MergeCells
        except ImportError:
            print("警告：openpyxl 未安装，无法生成横道图", file=sys.stderr)
            return None

        tasks = data_dict.get("tasks", [])
        if not tasks:
            print("警告：无任务数据，无法生成横道图", file=sys.stderr)
            return None

        if not output_path:
            timestamp = datetime.now().strftime("%Y%m%d%H%M%S")
            project_name = self._safe_filename(data_dict.get("project_name", "未命名工程"))
            output_path = os.path.join(self.output_dir, f"横道图_{project_name}_{timestamp}.xlsx")

        os.makedirs(os.path.dirname(output_path), exist_ok=True)

        # 计算日期范围
        from datetime import datetime as dt, timedelta
        all_starts = []
        all_ends = []
        for task in tasks:
            if task.get("start"):
                all_starts.append(dt.strptime(task["start"], "%Y-%m-%d"))
            if task.get("end"):
                all_ends.append(dt.strptime(task["end"], "%Y-%m-%d"))

        if not all_starts or not all_ends:
            return None

        date_start = min(all_starts)
        date_end = max(all_ends)

        # 生成日期列表
        dates = []
        current = date_start
        while current <= date_end:
            dates.append(current)
            current += timedelta(days=1)

        # Excel 日期基准：1899-12-30 = 序列号 1
        excel_base = dt(1899, 12, 30)

        # 样式
        bar_fill = PatternFill(start_color="00B0F0", end_color="00B0F0", fill_type="solid")
        month_header_font = Font(name="宋体", size=9, bold=True)
        day_header_font = Font(name="宋体", size=8)
        task_font = Font(name="宋体", size=10)
        center_align = Alignment(horizontal="center", vertical="center")
        left_align = Alignment(horizontal="left", vertical="center")
        thin_border = Border(
            left=Side(style="thin"),
            right=Side(style="thin"),
            top=Side(style="thin"),
            bottom=Side(style="thin"),
        )

        wb = openpyxl.Workbook()
        ws = wb.active
        ws.title = "施工进度横道图"

        # ---- 第 1 行：月份行（用跨列居中代替合并单元格，避免边框丢失） ----
        # 先按月份分组日期列
        month_groups = []
        prev_month = None
        for i, d in enumerate(dates):
            month_key = (d.year, d.month)
            col = i + 2  # B 列开始
            if month_key != prev_month:
                month_groups.append({
                    'label': f"{d.year}年{d.month}月",
                    'start_col': col,
                    'end_col': col,
                })
                prev_month = month_key
            else:
                month_groups[-1]['end_col'] = col

        header_fill = PatternFill(start_color="D9E1F2", end_color="D9E1F2", fill_type="solid")

        # A1: 任务名称
        a1_cell = ws.cell(row=1, column=1, value="任务名称")
        a1_cell.font = Font(name="宋体", size=9, bold=True)
        a1_cell.alignment = center_align
        a1_cell.border = thin_border
        a1_cell.fill = header_fill

        # 月份行：只在每月首列写月份标签，用跨列居中对齐
        for mg in month_groups:
            cell = ws.cell(row=1, column=mg['start_col'], value=mg['label'])
            cell.font = month_header_font
            cell.alignment = Alignment(horizontal="centerContinuous", vertical="center")
            cell.border = thin_border
            cell.fill = header_fill
            # 非首列的月份单元格也设置边框和填充
            for c in range(mg['start_col'] + 1, mg['end_col'] + 1):
                bc = ws.cell(row=1, column=c)
                bc.border = thin_border
                bc.fill = header_fill

        # ---- 第 2 行：日期行（显示日） ----
        a2_cell = ws.cell(row=2, column=1)
        a2_cell.border = thin_border
        a2_cell.fill = header_fill

        for i, d in enumerate(dates):
            col = i + 2
            serial = (d - excel_base).days
            cell = ws.cell(row=2, column=col, value=serial)
            cell.number_format = 'D'  # 只显示日
            cell.font = day_header_font
            cell.alignment = center_align
            cell.border = thin_border
            cell.fill = header_fill

        # ---- 第 3 行起：任务名称 + 横道条 ----
        # 先给所有数据单元格设置边框
        for ri, task in enumerate(tasks, 3):
            # A列任务名称
            name_cell = ws.cell(row=ri, column=1, value=task.get("name", ""))
            name_cell.font = task_font
            name_cell.alignment = left_align
            name_cell.border = thin_border

            # 所有日期列先加边框
            for col in range(2, len(dates) + 2):
                ws.cell(row=ri, column=col).border = thin_border

            # 绘制横道条：对应日期范围填充蓝色
            task_start = dt.strptime(task["start"], "%Y-%m-%d") if task.get("start") else None
            task_end = dt.strptime(task["end"], "%Y-%m-%d") if task.get("end") else None
            if task_start and task_end:
                for i, d in enumerate(dates):
                    if task_start <= d <= task_end:
                        col = i + 2
                        ws.cell(row=ri, column=col).fill = bar_fill

        # 设置列宽
        ws.column_dimensions["A"].width = 28  # 任务名称列加宽
        for i in range(len(dates)):
            ws.column_dimensions[get_column_letter(i + 2)].width = 4.5

        # 设置行高
        ws.row_dimensions[1].height = 22  # 月份行
        ws.row_dimensions[2].height = 18  # 日期行
        for ri in range(3, len(tasks) + 3):
            ws.row_dimensions[ri].height = 20

        # 冻结窗格：冻结前2行和A列
        ws.freeze_panes = "B3"

        # 打印设置：横向打印
        ws.page_setup.orientation = 'landscape'
        ws.page_setup.paperSize = ws.PAPERSIZE_A4
        ws.page_setup.fitToWidth = 1
        ws.page_setup.fitToHeight = 0
        ws.sheet_properties.pageSetUpPr.fitToPage = True

        wb.save(output_path)
        return output_path

    def _create_fallback_doc(self, data_dict):
        """创建兜底的基础格式文档

        当模板文件不存在或 .doc 转换失败时，生成一个包含所有数据的基础文档。

        Args:
            data_dict: 文档数据字典

        Returns:
            Document 对象
        """
        doc = Document()

        # 设置默认字体
        style = doc.styles["Normal"]
        font = style.font
        font.name = "宋体"
        font.size = Pt(12)

        # 文档标题
        title = data_dict.get("title", "施工资料文档")
        heading = doc.add_heading(title, level=0)
        heading.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # 基本信息
        doc.add_paragraph("")
        info_items = [
            ("工程名称", "project_name"),
            ("工程编号", "project_code"),
            ("分包单位", "subcontractor"),
            ("工程地点", "location"),
            ("编制日期", "date"),
        ]
        for label, key in info_items:
            value = data_dict.get(key, "")
            if value:
                p = doc.add_paragraph()
                p.add_run(f"{label}：").bold = True
                p.add_run(str(value))

        # 分隔线
        doc.add_paragraph("—" * 40)

        # 文档正文内容
        content = data_dict.get("content", "")
        if content:
            doc.add_heading("正文内容", level=1)
            if isinstance(content, list):
                for item in content:
                    doc.add_paragraph(str(item))
            else:
                doc.add_paragraph(str(content))

        # 人员信息
        workers = data_dict.get("workers", "")
        if workers:
            doc.add_heading("相关人员", level=1)
            if isinstance(workers, list):
                for w in workers:
                    doc.add_paragraph(str(w))
            else:
                doc.add_paragraph(str(workers))

        # 工艺信息
        processes = data_dict.get("processes", "")
        if processes:
            doc.add_heading("施工工艺", level=1)
            if isinstance(processes, list):
                for p in processes:
                    doc.add_paragraph(str(p))
            else:
                doc.add_paragraph(str(processes))

        return doc

    @staticmethod
    def _safe_filename(name, max_length=80):
        """将任意字符串转换为安全的文件名（替换 Windows 非法字符）"""
        if not name:
            return "未命名"
        invalid_chars = '<>:"/\\|?*'
        result = "".join("_" if c in invalid_chars else c for c in str(name))
        result = result.strip().rstrip(". ")
        if len(result) > max_length:
            result = result[:max_length]
        return result or "未命名"

    def generate_construction_design(self, data_dict):
        """生成施工组织设计文档

        Args:
            data_dict: 包含所有占位符数据的字典

        Returns:
            生成的文件路径
        """
        template_path = _find_template(self.template_dir, "施工组织设计模板")
        timestamp = datetime.now().strftime("%Y%m%d%H%M%S")
        project_name = self._safe_filename(data_dict.get("project_name", "未命名工程"))
        output_filename = f"施工组织设计_{project_name}_{timestamp}.docx"
        output_path = os.path.join(self.output_dir, output_filename)
        if template_path:
            return self.generate_from_template(template_path, output_path, data_dict)
        return self.generate_from_template(
            os.path.join(self.template_dir, "施工组织设计模板.docx"), output_path, data_dict
        )

    def generate_survey(self, data_dict):
        """生成项目勘察单文档

        Args:
            data_dict: 包含所有占位符数据的字典

        Returns:
            生成的文件路径
        """
        template_path = _find_template(self.template_dir, "项目勘察单模板")
        timestamp = datetime.now().strftime("%Y%m%d%H%M%S")
        project_name = self._safe_filename(data_dict.get("project_name", "未命名工程"))
        output_filename = f"项目勘察单_{project_name}_{timestamp}.docx"
        output_path = os.path.join(self.output_dir, output_filename)
        if template_path:
            return self.generate_from_template(template_path, output_path, data_dict)
        return self.generate_from_template(
            os.path.join(self.template_dir, "项目勘察单模板.docx"), output_path, data_dict
        )

    def generate_tech_briefing(self, data_dict):
        """生成技术交底文档

        Args:
            data_dict: 包含所有占位符数据的字典

        Returns:
            生成的文件路径
        """
        template_path = _find_template(self.template_dir, "技术交底模板")
        timestamp = datetime.now().strftime("%Y%m%d%H%M%S")
        project_name = self._safe_filename(data_dict.get("project_name", "未命名工程"))
        output_filename = f"技术交底_{project_name}_{timestamp}.docx"
        output_path = os.path.join(self.output_dir, output_filename)
        if template_path:
            return self.generate_from_template(template_path, output_path, data_dict)
        return self.generate_from_template(
            os.path.join(self.template_dir, "技术交底模板.docx"), output_path, data_dict
        )

    def generate_safety_briefing(self, data_dict):
        """生成安全交底文档

        Args:
            data_dict: 包含所有占位符数据的字典

        Returns:
            生成的文件路径
        """
        template_path = _find_template(self.template_dir, "安全交底模板")
        timestamp = datetime.now().strftime("%Y%m%d%H%M%S")
        project_name = self._safe_filename(data_dict.get("project_name", "未命名工程"))
        output_filename = f"安全交底_{project_name}_{timestamp}.docx"
        output_path = os.path.join(self.output_dir, output_filename)
        if template_path:
            return self.generate_from_template(template_path, output_path, data_dict)
        return self.generate_from_template(
            os.path.join(self.template_dir, "安全交底模板.docx"), output_path, data_dict
        )
